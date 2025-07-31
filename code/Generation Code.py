import json
from docx import Document
from google.colab import drive

# Mount Google Drive
drive.mount('/content/drive')

# Initialize 
key =
client =



def generate_sarcastic_conversation(incongruity, shock_value, context_dependency, emotion):

    prompt = f"""
You are a sarcasm simulation system. Create a short fictional dialogue that includes a clearly sarcastic utterance. Use the inputs below to guide the tone and structure.

Parameters:
- Incongruity Rating (1–10): {incongruity}
- Shock Value: {shock_value}
- Context Dependency: {context_dependency}
- Emotion of Sarcastic Utterance: {emotion}

Output format:

Conversation:
Speaker A: ...
Speaker B: ...
Speaker A: ...
(At least 3 turns)

Sarcastic Utterance: (copy the sarcastic utterance exactly here)
Sarcasm Type: (Self-deprecating, Brooding, Deadpan, Polite, Obnoxious, Raging, or Manic)
Emotion: {emotion}
Incongruity Rating: {incongruity}
Shock Value: {shock_value}
Context Dependency: {context_dependency}
"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()

# Document to store results
doc = Document()

# Define your different parameter sets to loop over
generation_parameters = [
    # SARCASTIC
    {"incongruity": 9, "shock_value": "moderate", "context_dependency": "medium", "emotion": "Surprise"},


]

# Loop through and generate
for idx, params in enumerate(generation_parameters, start=1):
    try:
        output = generate_sarcastic_conversation(**params)
        doc.add_heading(f"Example {idx}", level=1)
        for line in output.split("\n"):
            doc.add_paragraph(line.strip())
        doc.add_paragraph("\n" + "-"*50 + "\n")
    except Exception as e:
        print(f"Error generating Example {idx}: {e}")

# Save the document
save_path = '/content/drive/MyDrive/Algoverse research/generated_sarcasm_examples.docx'
doc.save(save_path)
print(f"Saved all generated examples to {save_path}")